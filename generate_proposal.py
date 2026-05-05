from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# --- Helper functions ---
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
    return h

def add_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '1B4332'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# --- HEADER ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('SERVICE PROPOSAL')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
run.font.bold = True
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
run = subtitle.add_run('POS System & Website Management')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.paragraph_format.space_after = Pt(6)
run = subtitle2.add_run('Prepared for Crisp on Creek')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_line()

# --- INFO TABLE ---
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
info_data = [
    ('Prepared by:', 'Matthias Campbell'),
    ('Date:', datetime.now().strftime('%d %B %Y')),
    ('Client:', 'Crisp on Creek — 1164 Cavendish Rd, Mt Gravatt East QLD'),
    ('Reference:', 'COC-2026-001'),
]
for i, (label, value) in enumerate(info_data):
    cell_l = info_table.cell(i, 0)
    cell_r = info_table.cell(i, 1)
    cell_l.text = label
    cell_r.text = value
    for cell in (cell_l, cell_r):
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    for run in cell_l.paragraphs[0].runs:
        run.font.bold = True

doc.add_paragraph()

# --- EXECUTIVE SUMMARY ---
add_heading_styled('Executive Summary', level=1)
doc.add_paragraph(
    'This proposal outlines a complete technology management package for Crisp on Creek, '
    'covering the custom-built Crisp POS point-of-sale system and the crisponcreek.com.au '
    'customer-facing website. Both systems will be maintained, updated, and supported on an '
    'ongoing basis for a fixed monthly fee of $350 (excl. GST).'
)
doc.add_paragraph(
    'This replaces the need for separate Profit Track licensing fees, third-party web hosting '
    'management, and ad-hoc IT support — consolidating everything into a single, predictable cost.'
)

# --- CURRENT STATE ---
add_heading_styled('Current Systems Overview', level=1)

add_heading_styled('Crisp on Creek Website', level=2)
doc.add_paragraph(
    'The existing website (crisponcreek.com.au) is a custom-built, mobile-responsive product '
    'catalogue and storefront featuring:'
)
bullets = [
    'Full product catalogue with 2,500+ items across 15 categories (Fruit, Vegetables, Deli, Grocery, Dairy, Bakery, etc.)',
    'Search with real-time predictive results',
    'Category browsing with subcategory filtering and sorting',
    'Product detail modals with images and pricing',
    'Shopping cart with favourites/wishlist system',
    'Store information: hours, location with embedded Google Map, contact form',
    'Mobile-optimised with responsive sidebar navigation',
    '"Top Picks This Week" promotional banner system',
    'Hosted on GitHub Pages (free hosting, reliable CDN)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading_styled('Crisp POS System', level=2)
doc.add_paragraph(
    'The Crisp POS system is a custom-built, offline-first point-of-sale application designed '
    'specifically for Crisp on Creek. It runs on the existing in-store server hardware and '
    'provides:'
)
bullets = [
    'Full register operation: barcode scanning, PLU lookup, department buttons, numpad entry',
    'Configurable keyboard layout editor matching the existing Profit Track button arrangement',
    'Cash, EFTPOS/card, and split payment processing',
    'Tax-inclusive pricing with automatic GST extraction (Australian standard)',
    'Hold/recall transactions, returns by receipt, supervisor overrides',
    'Deal engine: multi-buy, buy-X-get-Y-free, percentage/dollar discounts, combo deals',
    'Receipt printing via ESC/POS thermal printers',
    'Cash drawer management: float, pickup, drop, end-of-day reconciliation',
    'Z Report generation for daily close',
    'Staff PIN login with role-based access (cashier/manager/admin)',
    'Admin panel: products, staff, deals, specials, transactions, hardware, import, settings',
    'Built on Electron + SQLite — runs entirely offline, no internet dependency',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# --- PROPOSED INTEGRATION ---
add_heading_styled('Proposed API Integration', level=1)
doc.add_paragraph(
    'A key deliverable of this engagement is connecting the POS system to the website via a '
    'live API, so that product data, pricing, and specials flow automatically from the register '
    'to the online storefront.'
)

add_heading_styled('How It Works', level=2)
items = [
    ('Single Source of Truth',
     'Products are managed in the POS admin panel. When staff update a price, add a new product, '
     'or create a special — that change automatically appears on the website. No double-entry.'),
    ('Live Pricing & Specials',
     'Weekly specials set up in the POS "Deals" or "Specials" tab are reflected on the website\'s '
     '"Top Picks This Week" section in real-time.'),
    ('Stock Awareness',
     'If a product is discontinued or deactivated in the POS, it disappears from the website '
     'automatically.'),
    ('New Product Sync',
     'New products added at the register (via import or manual entry) become available on the '
     'website without any manual upload.'),
    ('Category & Image Management',
     'Product categories, images, and descriptions managed in one place flow to both the in-store '
     'system and the online catalogue.'),
]
for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ': ')
    run.font.bold = True
    p.add_run(desc)

# --- SCOPE OF SERVICE ---
add_heading_styled('Scope of Monthly Service', level=1)

add_heading_styled('POS System Management', level=2)
bullets = [
    'Ongoing software updates and bug fixes',
    'New feature development as requested (within reasonable scope)',
    'Keyboard layout changes and button configuration',
    'Product imports and bulk data updates',
    'Deal/special configuration assistance',
    'Hardware troubleshooting support (printer, scanner, drawer, scale)',
    'Database maintenance and backup management',
    'Multi-lane server setup and maintenance',
    'Staff training support for new features',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading_styled('Website Management', level=2)
bullets = [
    'Website hosting and uptime monitoring',
    'Content updates (banners, promotions, seasonal messaging)',
    'Product catalogue sync via POS API integration',
    'Weekly specials & "Top Picks" updates (automated from POS)',
    'Design updates and mobile optimisation',
    'SEO maintenance and Google Business integration',
    'Domain and DNS management',
    'Bug fixes and browser compatibility updates',
    'Analytics reporting (monthly traffic summary)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading_styled('API & Integration', level=2)
bullets = [
    'POS-to-website product sync API (build, maintain, monitor)',
    'Automated pricing updates from register to web',
    'Specials/deals sync to website promotional sections',
    'Future: online ordering integration (when ready)',
    'Future: delivery platform API connections (Uber Eats, DoorDash)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# --- PRICING ---
add_heading_styled('Pricing', level=1)

pricing_table = doc.add_table(rows=5, cols=2)
pricing_table.style = 'Light Shading Accent 1'
pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER

pricing_data = [
    ('Service', 'Monthly Cost'),
    ('POS System — maintenance, updates & support', '$200'),
    ('Website — hosting, updates & content management', '$100'),
    ('API Integration — sync, monitoring & maintenance', '$50'),
    ('Total', '$350 / month (excl. GST)'),
]
for i, (service, cost) in enumerate(pricing_data):
    pricing_table.cell(i, 0).text = service
    pricing_table.cell(i, 1).text = cost
    if i == 0 or i == 4:
        for cell in (pricing_table.cell(i, 0), pricing_table.cell(i, 1)):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('No lock-in contract. ')
run.font.bold = True
p.add_run(
    'Month-to-month billing. Either party may end the arrangement with 30 days\' written notice. '
    'All code, data, and content remain the property of Crisp on Creek at all times.'
)

# --- WHAT'S INCLUDED VS EXTRA ---
add_heading_styled('What\'s Included vs. Additional Work', level=1)

inc_table = doc.add_table(rows=7, cols=2)
inc_table.style = 'Light List Accent 1'
inc_data = [
    ('Included in $350/month', 'Quoted Separately'),
    ('Bug fixes & patches', 'Major new system builds (e.g. online ordering checkout)'),
    ('Minor feature additions', 'Third-party integrations (Xero, MYOB, Tyro)'),
    ('Product/price updates', 'Custom hardware procurement'),
    ('Layout & design changes', 'On-site visits (travel outside normal support)'),
    ('Weekly specials management', 'Training sessions beyond initial setup'),
    ('Server & backup maintenance', 'Data migration from other POS systems'),
]
for i, (inc, extra) in enumerate(inc_data):
    inc_table.cell(i, 0).text = inc
    inc_table.cell(i, 1).text = extra
    if i == 0:
        for cell in (inc_table.cell(i, 0), inc_table.cell(i, 1)):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

# --- TIMELINE ---
doc.add_paragraph()
add_heading_styled('Implementation Timeline', level=1)

timeline_table = doc.add_table(rows=5, cols=3)
timeline_table.style = 'Light Shading Accent 1'
tl_data = [
    ('Phase', 'Deliverable', 'Timeframe'),
    ('1', 'POS system deployed on existing server, all lanes operational', 'Week 1-2'),
    ('2', 'API integration built — products sync from POS to website', 'Week 2-3'),
    ('3', 'Website updated with live POS data feed, specials automation', 'Week 3-4'),
    ('4', 'Testing, staff training, go-live', 'Week 4'),
]
for i, row_data in enumerate(tl_data):
    for j, val in enumerate(row_data):
        timeline_table.cell(i, j).text = val
        if i == 0:
            for run in timeline_table.cell(i, j).paragraphs[0].runs:
                run.font.bold = True

# --- BENEFITS ---
doc.add_paragraph()
add_heading_styled('Benefits to Crisp on Creek', level=1)

benefits = [
    ('Replace Profit Track licensing',
     'No more annual POS software fees. Crisp POS is custom-built and owned by you.'),
    ('One point of contact',
     'No juggling between a POS vendor, a web developer, and an IT person. One call handles everything.'),
    ('Automatic price sync',
     'Update a price at the register and it appears on the website. No manual double-entry, no outdated web prices.'),
    ('Purpose-built for your shop',
     'The POS mirrors your existing Profit Track layout — same button positions, same workflow. '
     'Staff retraining is minimal.'),
    ('No internet dependency',
     'The POS runs offline. If the internet drops, the shop keeps operating. '
     'Sync happens automatically when connectivity returns.'),
    ('Future-ready',
     'Online ordering, delivery platform integration, and loyalty programs can be added '
     'incrementally without replacing the whole system.'),
]
for title_text, desc in benefits:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ' — ')
    run.font.bold = True
    p.add_run(desc)

# --- CONTACT ---
doc.add_paragraph()
add_line()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Prepared by Matthias Campbell')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
run.font.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run(datetime.now().strftime('%d %B %Y'))
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
output_path = r'C:\Users\Matthias\Downloads\mcpos\Crisp_on_Creek_Service_Proposal.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
